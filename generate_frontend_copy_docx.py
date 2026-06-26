#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成前端界面文案汇总 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_elm = tcPr.find(qn(f'w:{edge}'))
            if edge_elm is None:
                edge_elm = OxmlElement(f'w:{edge}')
                tcPr.append(edge_elm)
            edge_elm.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            edge_elm.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            edge_elm.set(qn('w:space'), '0')
            edge_elm.set(qn('w:color'), kwargs[edge].get('color', 'auto'))


def add_section(doc, title, rows):
    """添加一个带表格的章节"""
    doc.add_heading(title, level=2)
    if not rows:
        doc.add_paragraph("（本节暂无文案）")
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '文案'
    hdr_cells[1].text = '来源文件'
    hdr_cells[2].text = '说明'

    # 表头加粗
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)

    for text, source, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = text
        row_cells[1].text = source
        row_cells[2].text = desc
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # 节间距


def main():
    doc = Document()

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)

    # 标题
    title = doc.add_heading('YQSX 前端界面文案汇总', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(18)
        run.font.bold = True

    doc.add_paragraph('文档范围：shop-web-next/src 下所有页面与组件的用户可见文案')
    doc.add_paragraph()

    # 1. 全局与导航
    add_section(doc, '1. 全局与导航', [
        ('YQSX 智能零食商城', 'app/layout.tsx', '页面 title'),
        ('AI 驱动的智能购物体验 - 拍照识别零食，智能推荐商品', 'app/layout.tsx', '页面 description'),
        ('YQSX', 'components/layout/navbar.tsx', '品牌名'),
        ('智能零食', 'components/layout/navbar.tsx', '品牌标签'),
        ('首页 / 拍照识别 / DIN 推荐 / 商品列表 / 我的订单', 'components/layout/navbar.tsx', '导航项'),
        ('登录', 'components/layout/navbar.tsx', '未登录按钮'),
        ('我的订单 / 退出登录', 'components/layout/navbar.tsx', '用户下拉菜单'),
        ('ID: {userInfo.userId}', 'components/layout/navbar.tsx', '用户 ID 展示'),
        ('页面出错了 / 抱歉，页面加载时遇到问题。可以尝试重新加载，或返回首页继续浏览。 / 重试 / 返回首页', 'app/error.tsx', '路由错误边界'),
        ('应用发生严重错误 / 抱歉，应用遇到致命问题。请尝试刷新页面，若问题持续请联系技术支持。 / 重试 / 刷新页面', 'app/global-error.tsx', '全局错误边界'),
        ('双击我可以打开智能问答哦~ / 点我一下，我会跳舞哦~ / 去拍照识别看看新零食吧！ / 今天想吃什么口味的零食？ / 我可以陪你逛一整天~ / 试试双击我的脑袋！', 'components/fun/snack-mascot.tsx', '吉祥物悬停提示'),
        ('双击打开问答~ / 小人在睡觉... / 双击我打开智能问答 ~', 'components/fun/snack-mascot.tsx', '吉祥物状态提示'),
        ('零食小助手', 'components/fun/snack-mascot.tsx', 'aria-label'),
    ])

    # 2. 首页
    add_section(doc, '2. 首页', [
        ('AI 智能 / 零食商城', 'app/page.tsx', 'Hero 标题'),
        ('拍照识别零食，AI 智能推荐，一键下单，享受未来购物体验。', 'app/page.tsx', 'Hero 描述'),
        ('开始识别 / 浏览商品', 'app/page.tsx', 'Hero 按钮'),
        ('AI 驱动的智能购物体验', 'components/ui/pixel-hero.tsx', 'Hero 顶部标签'),
        ('核心能力', 'app/page.tsx', '功能区标签'),
        ('让零食购物更简单', 'app/page.tsx', '功能区标题'),
        ('结合 DIN 推荐算法与智能前端展示，打造从登录、推荐到结果展示的一站式体验。', 'app/page.tsx', '功能区描述'),
        ('智能识别 / DIN 推荐 / 便捷下单', 'app/page.tsx', '功能卡片标题'),
        ('拍照即可识别零食，精准度高达95%，支持多种零食品类同时识别。 / 输入用户 ID 即可返回 Top40 纯数字商品 ID，直连推荐服务。 / 一键购买，快速完成订单，享受流畅的购物体验。', 'app/page.tsx', '功能卡片描述'),
        ('了解更多', 'app/page.tsx', '卡片悬停链接'),
        ('19+ 零食品类 / 95% 识别精度 / 3s 秒级响应', 'app/page.tsx', '数据展示区'),
        ('AI 智能识别', 'app/page.tsx', 'CTA 标签'),
        ('准备好发现新口味了吗？', 'app/page.tsx', 'CTA 标题'),
        ('立即体验 DIN 推荐结果展示，让每一位用户都能看到自己的 Top40 商品 ID。', 'app/page.tsx', 'CTA 描述'),
        ('立即体验', 'app/page.tsx', 'CTA 按钮'),
        ('© 2026 YQSX 智能零食商城. All rights reserved.', 'app/page.tsx', 'Footer 版权'),
        ('AI 驱动 · 智能识别 · 极速下单', 'app/page.tsx', 'Footer 标语'),
    ])

    # 3. 登录/注册/重置密码
    add_section(doc, '3. 登录 / 注册 / 重置密码', [
        ('登录失败，请检查用户名和密码', 'app/login/page.tsx', '登录失败提示'),
        ('智能零食商城', 'components/ui/auth-fuse.tsx', '左侧品牌标题'),
        ('发现每一口零食的惊喜，AI 让购物更有趣。', 'components/ui/auth-fuse.tsx', '左侧标语'),
        ('智能识别 / AI 推荐 / 便捷下单 / 极速响应', 'components/ui/auth-fuse.tsx', '左侧功能点'),
        ('登录 / 注册 / 重置密码', 'components/ui/auth-fuse.tsx', '标签页'),
        ('欢迎回来 / 登录您的账户，开启智能零食之旅', 'components/ui/auth-fuse.tsx', '登录模式标题'),
        ('创建账户 / 手机号验证注册，安全又便捷', 'components/ui/auth-fuse.tsx', '注册模式标题'),
        ('重置密码 / 验证手机号后设置新密码', 'components/ui/auth-fuse.tsx', '重置模式标题'),
        ('用户名 / 请输入用户名', 'components/auth/login-form.tsx', '登录表单'),
        ('密码 / 请输入密码', 'components/auth/login-form.tsx', '登录表单'),
        ('记住我 / 忘记密码？', 'components/auth/login-form.tsx', '登录辅助'),
        ('登录 / 处理中...', 'components/auth/login-form.tsx', '登录按钮'),
        ('手机号 / 请输入 11 位手机号', 'components/auth/register-form.tsx', '注册表单'),
        ('验证码 / 请输入 6 位验证码 / 获取验证码 / {countdown}s 后重发', 'components/auth/register-form.tsx', '注册表单'),
        ('演示环境验证码将输出在后端控制台', 'components/auth/register-form.tsx', '注册提示'),
        ('昵称（可选） / 请输入昵称', 'components/auth/register-form.tsx', '注册表单'),
        ('密码（至少 6 位） / 请再次输入密码', 'components/auth/register-form.tsx', '注册表单'),
        ('注册 / 处理中... / 返回登录', 'components/auth/register-form.tsx', '注册按钮'),
        ('新密码（至少 6 位） / 请再次输入新密码', 'components/auth/reset-form.tsx', '重置表单'),
        ('重置密码 / 处理中... / 返回登录', 'components/auth/reset-form.tsx', '重置按钮'),
        ('请输入有效的 11 位手机号 / 请输入 6 位数字验证码 / 密码长度不能少于 6 位 / 两次输入的密码不一致', 'components/ui/auth-fuse.tsx', '表单校验'),
        ('验证码已发送，请注意查收 / 演示验证码 / 已为您生成验证码，请输入后完成验证 / 知道了', 'components/ui/auth-fuse.tsx', '验证码弹窗'),
        ('测试账号：admin / 123456（仅开发环境）', 'components/ui/auth-fuse.tsx', '开发环境提示'),
        ('显示密码 / 隐藏密码', 'components/ui/auth-primitives.tsx', '密码框 aria-label'),
    ])

    # 4. 商品列表
    add_section(doc, '4. 商品列表', [
        ('为您推荐', 'app/products/page.tsx', '推荐区标题'),
        ('正在加载推荐结果... / 推荐请求被限流，请稍后再试 / 基于用户 {id} · 已降级为热销商品 · {latency}ms / 基于用户 {id} 的个性化推荐 · {缓存状态} · {latency}ms', 'app/products/page.tsx', '推荐状态'),
        ('换一批 / 加载中...', 'app/products/page.tsx', '换一批按钮'),
        ('暂无推荐', 'app/products/page.tsx', '推荐空状态'),
        ('排名 {index + 1} / ¥{product.price} / ID: {product.id}', 'app/products/page.tsx', '推荐卡片'),
        ('精选好物', 'app/products/page.tsx', '商品区标签'),
        ('发现美味零食', 'app/products/page.tsx', '商品区标题'),
        ('精选 {products.length} 款美味零食，每一口都是惊喜', 'app/products/page.tsx', '商品区描述'),
        ('全部 / {category}', 'app/products/page.tsx', '分类筛选'),
        ('暂无商品 / 去拍照识别', 'app/products/page.tsx', '商品空状态'),
        ('¥{product.price} / 库存 {product.stock} / 销量: {product.sales} / 查看详情', 'app/products/page.tsx', '商品卡片'),
    ])

    # 5. 商品详情
    add_section(doc, '5. 商品详情', [
        ('加载商品失败，请刷新重试', 'app/products/[id]/page.tsx', '加载失败'),
        ('商品不存在 / 返回商品列表', 'app/products/[id]/page.tsx', '空状态'),
        ('返回', 'app/products/[id]/page.tsx', '返回按钮'),
        ('商品详情', 'app/products/[id]/page.tsx', '详情标签'),
        ('{product.categoryName} / 零食', 'app/products/[id]/page.tsx', '分类徽章'),
        ('4.8 分', 'app/products/[id]/page.tsx', '评分'),
        ('售价 / ¥{product.price}', 'app/products/[id]/page.tsx', '价格'),
        ('库存: {product.stock} 件 / 已售: {product.sales} 件', 'app/products/[id]/page.tsx', '库存销量'),
        ('商品介绍 / {product.name} 是一款优质零食，精选原料，口感醇厚。无论是自己享用还是分享给朋友，都是绝佳的选择。', 'app/products/[id]/page.tsx', '商品介绍'),
        ('立即购买 / 已售罄 / 创建订单中...', 'app/products/[id]/page.tsx', '购买按钮'),
        ('登录后可立即购买，去登录', 'app/products/[id]/page.tsx', '未登录提示'),
        ('创建订单失败 / 创建订单失败，请重试', 'app/products/[id]/page.tsx', '下单失败 alert'),
    ])

    # 6. 订单
    add_section(doc, '6. 订单', [
        ('登录后查看订单 / 去登录', 'app/orders/page.tsx', '未登录空状态'),
        ('订单中心 / 我的订单 / 共 {orders.length} 笔订单，随时查看购买记录', 'app/orders/page.tsx', '订单列表标题区'),
        ('全部 / 待支付 / 已支付 / 已完成 / 已取消', 'app/orders/page.tsx', '状态筛选'),
        ('暂无订单 / 该状态下暂无订单 / 去购物', 'app/orders/page.tsx', '订单空状态'),
        ('订单编号 / #{order.id}', 'app/orders/page.tsx', '订单卡片'),
        ('数量: {getOrderQuantity(order)} / ¥{amount} / 查看详情', 'app/orders/page.tsx', '订单卡片'),
        ('待支付 / 已支付 / 已完成 / 已取消', 'components/orders/order-status-badge.tsx', '状态徽章'),
        ('加载订单失败，请刷新重试', 'app/orders/[id]/page.tsx', '加载失败'),
        ('订单不存在 / 返回订单列表', 'app/orders/[id]/page.tsx', '空状态'),
        ('返回', 'app/orders/[id]/page.tsx', '返回按钮'),
        ('订单详情 / {message.title}', 'app/orders/[id]/page.tsx', '详情标题'),
        ('订单待支付 / 请在 30 分钟内完成支付，超时订单将自动取消', 'app/orders/[id]/page.tsx', '待支付状态'),
        ('订单已支付 / 商家正在准备发货，请耐心等待', 'app/orders/[id]/page.tsx', '已支付状态'),
        ('订单已完成 / 感谢你的购买，期待再次光临', 'app/orders/[id]/page.tsx', '已完成状态'),
        ('订单已取消 / 该订单已被取消，如有疑问请联系客服', 'app/orders/[id]/page.tsx', '已取消状态'),
        ('订单编号 / 下单用户 / {order.username}', 'app/orders/[id]/page.tsx', '订单信息'),
        ('商品信息 / 数量: {order.number} 件 / 单价: ¥{order.productPrice} / 订单金额 / ¥{totalAmount}', 'app/orders/[id]/page.tsx', '商品信息'),
        ('立即支付', 'app/orders/[id]/page.tsx', '支付按钮'),
    ])

    # 7. 支付
    add_section(doc, '7. 支付', [
        ('收银台 / 订单支付', 'components/payment/idle-view.tsx', '支付页标题'),
        ('订单编号: #{order.id} / {productName} x {number} / 支付金额 / ¥{totalAmount}', 'components/payment/idle-view.tsx', '订单信息'),
        ('剩余支付时间 / {formatTime(countdown)}', 'components/payment/idle-view.tsx', '倒计时'),
        ('确认支付 / 支付已超时', 'components/payment/idle-view.tsx', '支付按钮'),
        ('点击确认支付即表示你同意相关服务协议', 'components/payment/idle-view.tsx', '协议提示'),
        ('支付处理中... / 请稍候，不要关闭页面', 'components/payment/processing-view.tsx', '处理中'),
        ('支付成功！ / 正在跳转到订单详情... / 立即查看订单', 'components/payment/success-view.tsx', '成功页'),
        ('支付失败 / 支付超时，请重新下单 / 支付失败，请重试 / 返回商品列表 / 重新支付', 'components/payment/failed-view.tsx', '失败页'),
        ('加载订单失败，请刷新重试 / 重新加载', 'app/payment/[orderId]/page.tsx', '加载失败'),
    ])

    # 8. 拍照识别
    add_section(doc, '8. 拍照识别', [
        ('AI 智能识别 / 拍照识别零食 / 上传一张照片，AI instantly 识别图中的零食，并为你推荐相似美味', 'app/recognize/page.tsx', '页面标题'),
        ('上传照片', 'app/recognize/page.tsx', '上传区标题'),
        ('识别结果', 'app/recognize/page.tsx', '结果区标题'),
        ('DIN 推荐商品 ID / 基于当前登录用户的历史行为生成 Top40 纯数字商品 ID。 / {count} 条推荐 / 加载中', 'app/recognize/page.tsx', '推荐区标题'),
        ('当前暂无推荐数据，请先完成登录并确保算法服务可访问。', 'app/recognize/page.tsx', '推荐空状态'),
        ('推荐服务暂时不可用', 'app/recognize/page.tsx', '推荐错误'),
        ('点击或拖拽上传图片 / 支持 JPG、PNG 格式 / 也可以直接拍照上传', 'components/recognize/image-uploader.tsx', '上传组件'),
        ('待识别图片', 'components/recognize/image-uploader.tsx', '图片 alt'),
        ('识别失败 / {error} / 重新上传', 'components/recognize/result-panel.tsx', '识别错误'),
        ('检测到的零食 / {detectedCount} 个目标', 'components/recognize/result-panel.tsx', '检测结果'),
        ('未检测到零食，请尝试上传更清晰的图片', 'components/recognize/result-panel.tsx', '无目标提示'),
        ('为你推荐 / 共 {products.length} 件 / 浏览全部商品', 'components/recognize/result-panel.tsx', '识别推荐'),
        ('库存: {product.stock} / ¥{product.price}', 'components/recognize/result-panel.tsx', '推荐商品信息'),
        ('上传图片后，识别结果将显示在这里', 'components/recognize/result-panel.tsx', '空状态'),
        ('请上传图片文件 / 识别服务暂时不可用，请稍后重试', 'hooks/use-recognition.ts', '错误提示'),
    ])

    # 9. KBOQA 智能问答
    add_section(doc, '9. KBOQA 智能问答', [
        ('KBOQA 智能商品问答 / AI 智能商品问答，专业解答购物问题 / 融合电商领域知识库、大模型推理与智能推荐，为每一次购物决策提供可信答案', 'app/kboqa-chat/page.tsx', '页面标题'),
        ('消息历史 / {messages.length} 条消息', 'app/kboqa-chat/page.tsx', '聊天面板标题'),
        ('已收到你的问题，正在调用智能引擎分析...', 'app/kboqa-chat/page.tsx', '系统处理中消息'),
        ('KBOQA 处理完成 (置信度: {x}%) · 来源: {sources}', 'app/kboqa-chat/page.tsx', '系统完成消息'),
        ('选择下方模板或输入问题开始问答 / 同学A · {time} / 同学B (KBOQA) · {time}', 'app/kboqa-chat/page.tsx', '消息气泡'),
        ('推荐商品： / {category} · ¥{price} · {reason}', 'app/kboqa-chat/page.tsx', '推荐商品'),
        ('提交', 'app/kboqa-chat/page.tsx', '提交按钮'),
        ('输入商品问题，例如：商品 Littleswan/小天鹅 TH100-HL02T 是洗衣机吗？', 'app/kboqa-chat/page.tsx', '输入框占位符'),
        ('KBOQA 状态', 'app/kboqa-chat/page.tsx', '状态面板标题'),
        ('空闲待命 / 接收问题中 / 分析问题类型 / 生成回答中 / 整理推荐结果 / 完成', 'lib/kboqa/orchestrator.ts', '状态标签'),
        ('处理流程 / K 知识检索 / B 规则判定 / O 组装上下文 / QA 生成答案', 'app/kboqa-chat/page.tsx', '流程说明'),
        ('最近结果 / 答案类型 / 置信度 / 推荐数 / Trace ID / 数据来源', 'app/kboqa-chat/page.tsx', '结果面板'),
        ('品类判断 / 品牌知识 / 商品推荐 / 商品链接', 'lib/kboqa/templates.ts', '模板标签'),
        ('商品 Littleswan/小天鹅... / BAI/菜百首饰... / 帮我推荐机油滤芯... / 查看 MANN FILTER...', 'lib/kboqa/templates.ts', '模板示例问题'),
        ('判断商品是否属于某个细粒度品类 / 查询品牌的基本信息和主营类目 / 根据品类或需求推荐相关商品 / 获取商品的详情页链接', 'lib/kboqa/templates.ts', '模板描述'),
        ('KBOQA 智能商品问答 / 双击小人打开问答，选择模板或直接输入问题 / 预输入模板 / 进入完整问答界面 / 关闭', 'components/kboqa/qa-dialog.tsx', '悬浮问答窗'),
        ('KBOQA 请求失败 ({status})', 'lib/kboqa/orchestrator.ts', '请求错误'),
    ])

    # 10. 通用状态组件
    add_section(doc, '10. 通用状态组件', [
        ('{message} / 重新加载', 'components/async-state/error-state.tsx', '错误状态'),
        ('{title} / {description}', 'components/async-state/empty-state.tsx', '空状态'),
    ])

    # 11. API / 全局错误提示
    add_section(doc, '11. API / 全局错误提示', [
        ('登录已过期，请重新登录', 'lib/api/request.ts', '401 提示'),
        ('无权限访问该资源', 'lib/api/request.ts', '403 提示'),
        ('请求过于频繁，请稍后再试', 'lib/api/request.ts', '429 提示'),
        ('服务暂时不可用，请稍后再试', 'lib/api/request.ts', '503 提示'),
        ('无法连接到服务网关，请确认 Gateway 已启动', 'lib/api/request.ts', '无响应提示'),
        ('注册失败，请稍后再试 / 密码重置失败，请稍后再试 / 验证码发送失败，请稍后再试', 'lib/stores/auth.ts', 'auth store 错误兜底'),
    ])

    # 12. 零食类别名称
    add_section(doc, '12. 零食类别名称（识别展示用）', [
        ('Ashi Mashi 零食 / Chee pellet 番茄酱味 / Chee pellet 醋味 / Cheetoz 辣椒味薯片 / Cheetoz 番茄酱味薯片 / Cheetoz 洋葱香菜味薯片 / Cheetoz 咸味薯片 / Cheetoz 零食 30g / Cheetoz 零食 90g / Cheetoz 醋味薯片 / Cheetoz 车轮零食 / Maz Maz 番茄酱味薯片 / Maz Maz 土豆条 / Maz Maz 咸味薯片 / Maz Maz 醋味薯片 / Mini Lina 饼干 / Minoo 奶油饼干 / Naderi 迷你曲奇 / Naderi 迷你威化', 'lib/utils/category-mapping-data.ts', '识别结果展示名'),
        ('混合零食 / 膨化食品 / 薯片 / 薯条 / 饼干 / 威化', 'lib/utils/category-mapping-data.ts', '分类标签'),
    ])

    output_path = r'd:\Programming\YQSX\docs\前端界面文案汇总.docx'
    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')


if __name__ == '__main__':
    main()
