"""
生成三份不同排版风格的环境安装报告（Word格式）
模型层配置 - YOLO11 GPU环境搭建
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_DIR = r"d:\Programming\YQSX\docs"

# ============================================================
# 公共数据
# ============================================================
REPORT_DATE = datetime.date.today().strftime("%Y年%m月%d日")
REPORT_DATE_SHORT = datetime.date.today().strftime("%Y-%m-%d")

CONDA_PACKAGES = [
    ("matplotlib", ">=3.7.0", "数据可视化与图表绘制"),
    ("numpy", ">=1.24.0", "高性能数值计算库"),
    ("pandas", ">=2.0.0", "结构化数据处理与分析"),
    ("Pillow", ">=10.0.0", "图像处理与格式支持"),
    ("PyYAML", ">=6.0", "YAML配置文件解析"),
    ("scikit-learn", ">=1.3.0", "机器学习算法工具包"),
    ("scipy", ">=1.10.0", "科学计算与工程算法"),
    ("torch", ">=2.0.0", "PyTorch深度学习框架核心"),
    ("torchaudio", ">=2.0.0", "PyTorch音频处理模块"),
    ("torchvision", ">=0.15.0", "PyTorch计算机视觉模块"),
    ("tqdm", ">=4.65.0", "进度条显示工具"),
]

EXTRA_MODEL_PACKAGES = [
    ("ultralytics", ">=8.0.0", "YOLO系列模型框架（YOLO11）"),
    ("opencv-python", ">=4.8.0", "计算机视觉图像处理库"),
    ("imagehash", ">=4.3.0", "图像感知哈希与去重"),
    ("seaborn", ">=0.12.0", "统计数据分析可视化"),
    ("loguru", ">=0.7.0", "高性能日志记录工具"),
]

YOLO_CONFIGS = [
    ("baseline_config", "yolo11n.pt", "50", "16", "640", "快速验证与基准测试"),
    ("fast_iteration_config", "yolo11n.pt", "30", "16", "640", "快速迭代实验"),
    ("optimized_config", "yolo11s.pt", "100", "16", "640", "精度与速度平衡（推荐）"),
    ("high_precision_config", "yolo11m.pt", "150", "12", "640", "最高精度场景"),
]

DATASET_INFO = {
    "classes": 19,
    "source": "RoboFlow - Iranian Snack and Chips Detection",
    "train_path": "train/images",
    "val_path": "valid/images",
    "test_path": "test/images",
    "license": "CC BY 4.0",
}


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ============================================================
# 报告一：传统企业风格
# ============================================================
def generate_report_1():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ---- 封面 ----
    for _ in range(6):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("模型层环境安装配置报告")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph("")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("—— 基于YOLO11的GPU检测环境搭建")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    for _ in range(4):
        doc.add_paragraph("")

    # 封面信息表
    cover_info = [
        ("项目名称", "零食识别系统 - 模型层环境配置"),
        ("报告编号", "ENV-RPT-2026-001"),
        ("版    本", "V1.0"),
        ("编制日期", REPORT_DATE),
        ("密    级", "内部"),
    ]
    cover_table = doc.add_table(rows=len(cover_info), cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(cover_info):
        cell_k = cover_table.rows[i].cells[0]
        cell_v = cover_table.rows[i].cells[1]
        cell_k.width = Cm(4)
        cell_v.width = Cm(10)

        cell_k.text = ""
        p = cell_k.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(k)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        set_cell_shading(cell_k, "D6E4F0")

        cell_v.text = ""
        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(v)
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    doc.add_page_break()

    # ---- 目录页 ----
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目    录")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph("")

    toc_items = [
        ("第一章  概述", 1),
        ("  1.1 编写目的", 2),
        ("  1.2 适用范围", 2),
        ("  1.3 术语说明", 2),
        ("第二章  环境要求", 1),
        ("  2.1 硬件要求", 2),
        ("  2.2 软件要求", 2),
        ("  2.3 GPU驱动要求", 2),
        ("第三章  环境安装步骤", 1),
        ("  3.1 GPU驱动安装", 2),
        ("  3.2 Conda环境管理工具安装", 2),
        ("  3.3 虚拟环境创建", 2),
        ("  3.4 依赖包安装", 2),
        ("  3.5 YOLO11模型框架安装", 2),
        ("  3.6 开发工具配置", 2),
        ("第四章  模型层配置", 1),
        ("  4.1 YOLO11模型说明", 2),
        ("  4.2 训练配置参数", 2),
        ("  4.3 数据集配置", 2),
        ("第五章  环境验证", 1),
        ("  5.1 验证步骤", 2),
        ("  5.2 验证结果", 2),
        ("第六章  附录", 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        if level == 1:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        else:
            run.font.size = Pt(12)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---- 正文 ----

    # 辅助函数
    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return h

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        return p

    # 第一章
    add_heading_styled("第一章  概述", 1)

    add_heading_styled("1.1 编写目的", 2)
    add_body("本报告旨在记录模型层环境的完整安装与配置过程，涵盖GPU驱动、Conda环境管理工具、Python虚拟环境、深度学习框架及YOLO11模型框架的安装步骤，为后续环境复现与运维提供标准化参考文档。")

    add_heading_styled("1.2 适用范围", 2)
    add_body("本报告适用于零食识别系统模型层的开发环境搭建，目标读者为参与模型训练与推理的开发工程师及测试工程师。")

    add_heading_styled("1.3 术语说明", 2)
    terms = [
        ("Conda", "开源的包管理与环境管理系统，支持Python及非Python包的安装与隔离"),
        ("YOLO11", "Ultralytics公司发布的最新目标检测模型，支持n/s/m/l/x多种规格"),
        ("CUDA", "NVIDIA推出的并行计算平台，用于GPU加速计算"),
        ("cuDNN", "NVIDIA深度神经网络加速库，为卷积等操作提供高度优化的实现"),
        ("ultralytics", "YOLO系列模型的官方Python框架，提供训练、验证、推理一体化接口"),
    ]
    add_formatted_table(doc, ["术语", "说明"], terms, [3, 11])

    doc.add_paragraph("")

    # 第二章
    add_heading_styled("第二章  环境要求", 1)

    add_heading_styled("2.1 硬件要求", 2)
    hw_rows = [
        ("CPU", "Intel Core i5 及以上 / AMD Ryzen 5 及以上"),
        ("内存", "16GB 及以上（推荐32GB）"),
        ("硬盘", "SSD，可用空间 >= 50GB"),
        ("GPU", "NVIDIA显卡，显存 >= 4GB（推荐 >= 8GB）"),
    ]
    add_formatted_table(doc, ["项目", "最低要求"], hw_rows, [3, 11])

    doc.add_paragraph("")

    add_heading_styled("2.2 软件要求", 2)
    sw_rows = [
        ("操作系统", "Windows 10/11 64位 或 Ubuntu 20.04+"),
        ("NVIDIA驱动", ">= 525.60（支持CUDA 12.x）"),
        ("CUDA Toolkit", "12.1 及以上"),
        ("cuDNN", "8.9 及以上"),
        ("Conda", "Anaconda3 或 Miniconda3（最新版）"),
        ("Python", "3.7"),
        ("PyCharm", "Professional 2023.3+ （推荐）/ Community"),
        ("Jupyter", "jupyterlab >= 4.0"),
    ]
    add_formatted_table(doc, ["软件", "版本要求"], sw_rows, [4, 10])

    doc.add_paragraph("")

    add_heading_styled("2.3 GPU驱动要求", 2)
    add_body("安装GPU驱动后，需确认CUDA Toolkit与cuDNN版本兼容。推荐使用nvidia-smi命令验证驱动状态，确保GPU设备被系统正确识别。")

    doc.add_page_break()

    # 第三章
    add_heading_styled("第三章  环境安装步骤", 1)

    add_heading_styled("3.1 GPU驱动安装", 2)
    add_body("步骤一：访问NVIDIA官网下载对应型号的最新驱动程序。")
    add_body('步骤二：运行安装程序，选择"自定义安装"，勾选"执行清洁安装"选项。')
    add_body("步骤三：安装完成后，打开命令提示符执行以下命令验证：")
    p = doc.add_paragraph()
    run = p.add_run("    nvidia-smi")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    add_body("步骤四：确认输出中显示GPU型号、驱动版本及CUDA版本信息。")

    add_heading_styled("3.2 Conda环境管理工具安装", 2)
    add_body("步骤一：下载Anaconda3安装包（完整版，含env路径管理功能）。")
    add_body("步骤二：运行安装程序，安装路径建议为非系统盘，如 D:\\Anaconda3。")
    add_body('步骤三：安装过程中勾选"Add Anaconda to PATH"（或手动配置环境变量）。')
    add_body("步骤四：安装完成后打开Anaconda Prompt，执行以下命令确认：")
    p = doc.add_paragraph()
    run = p.add_run("    conda --version")
    run.font.name = "Consolas"
    run.font.size = Pt(11)

    add_heading_styled("3.3 虚拟环境创建", 2)
    add_body("使用Conda创建名为AIDetection的虚拟环境，指定Python版本为3.7：")
    p = doc.add_paragraph()
    run = p.add_run("    conda create -n AIDetection python=3.7")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True
    add_body("激活环境：")
    p = doc.add_paragraph()
    run = p.add_run("    conda activate AIDetection")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    add_body("环境路径默认位于 Anaconda安装目录\\envs\\AIDetection。")

    add_heading_styled("3.4 依赖包安装", 2)
    add_body("在激活的AIDetection环境中，执行以下命令安装核心依赖包：")
    p = doc.add_paragraph()
    cmd = "    pip install matplotlib numpy pandas Pillow PyYAML scikit-learn scipy torch torchaudio torchvision tqdm"
    run = p.add_run(cmd)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    add_body("各依赖包详细信息如下表所示：")
    pkg_rows = [(name, ver, desc) for name, ver, desc in CONDA_PACKAGES]
    add_formatted_table(doc, ["包名", "版本要求", "用途说明"], pkg_rows, [3.5, 3, 8])

    doc.add_paragraph("")
    add_body("模型层额外依赖包（项目中已使用）：")
    extra_rows = [(name, ver, desc) for name, ver, desc in EXTRA_MODEL_PACKAGES]
    add_formatted_table(doc, ["包名", "版本要求", "用途说明"], extra_rows, [3.5, 3, 8])

    doc.add_paragraph("")

    add_heading_styled("3.5 YOLO11模型框架安装", 2)
    add_body("YOLO11模型通过ultralytics框架提供，安装命令如下：")
    p = doc.add_paragraph()
    run = p.add_run("    pip install ultralytics>=8.0.0")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True
    add_body("安装完成后，验证安装是否成功：")
    p = doc.add_paragraph()
    run = p.add_run("    python -c \"from ultralytics import YOLO; print('ultralytics OK')\"")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    add_heading_styled("3.6 开发工具配置", 2)
    add_body("（1）PyCharm配置（推荐）")
    add_body("步骤一：安装PyCharm Professional版。")
    add_body("步骤二：打开项目后，进入 File → Settings → Project → Python Interpreter。")
    add_body("步骤三：点击齿轮图标 → Add Interpreter → Conda Environment → Existing environment。")
    add_body("步骤四：选择AIDetection环境的Python解释器路径。")
    add_body("步骤五：确认Interpreter路径指向 Anaconda\\envs\\AIDetection\\python.exe。")
    add_body("")
    add_body("（2）Jupyter配置")
    add_body("安装JupyterLab：")
    p = doc.add_paragraph()
    run = p.add_run("    pip install jupyterlab>=4.0")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    add_body("安装ipykernel以支持Conda环境：")
    p = doc.add_paragraph()
    run = p.add_run("    pip install ipykernel\n    python -m ipykernel install --user --name AIDetection --display-name \"Python (AIDetection)\"")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    add_body("启动JupyterLab：")
    p = doc.add_paragraph()
    run = p.add_run("    jupyter lab")
    run.font.name = "Consolas"
    run.font.size = Pt(11)

    doc.add_page_break()

    # 第四章
    add_heading_styled("第四章  模型层配置", 1)

    add_heading_styled("4.1 YOLO11模型说明", 2)
    add_body("本项目采用YOLO11作为目标检测模型，YOLO11是Ultralytics发布的最新一代目标检测架构，相较于前代在精度与速度上均有显著提升。项目中使用以下三种规格模型：")

    model_rows = [
        ("yolo11n.pt", "Nano", "参数量最小，推理速度最快，适用于快速验证"),
        ("yolo11s.pt", "Small", "精度与速度的最佳平衡，推荐作为默认配置"),
        ("yolo11m.pt", "Medium", "参数量较大，精度最高，适用于对延迟不敏感的场景"),
    ]
    add_formatted_table(doc, ["模型文件", "规格", "特点"], model_rows, [3, 2, 9.5])

    doc.add_paragraph("")

    add_heading_styled("4.2 训练配置参数", 2)
    add_body("项目提供四套训练配置文件，分别对应不同的训练策略：")

    config_rows = [(c[0], c[1], c[2], c[3], c[4], c[5]) for c in YOLO_CONFIGS]
    add_formatted_table(doc,
        ["配置文件", "模型", "Epochs", "Batch", "ImgSz", "适用场景"],
        config_rows, [3.5, 2.2, 1.5, 1.3, 1.5, 4.5])

    doc.add_paragraph("")
    add_body("通用训练参数说明：")
    train_params = [
        ("optimizer", "AdamW", "优化器算法"),
        ("lr0", "0.001", "初始学习率"),
        ("lrf", "0.01", "最终学习率系数"),
        ("patience", "15~30", "早停耐心值"),
        ("cos_lr", "True", "余弦退火学习率调度"),
        ("mosaic", "0.8~1.0", "Mosaic数据增强概率"),
        ("mixup", "0.0~0.2", "MixUp数据增强概率"),
        ("label_smoothing", "0.0~0.1", "标签平滑系数"),
        ("amp", "True", "自动混合精度训练"),
    ]
    add_formatted_table(doc, ["参数", "取值", "说明"], train_params, [3.5, 3, 8])

    doc.add_paragraph("")

    add_heading_styled("4.3 数据集配置", 2)
    add_body("数据集配置文件（data.yaml）定义了训练、验证、测试数据路径及类别信息：")

    ds_rows = [
        ("数据集来源", DATASET_INFO["source"]),
        ("类别数量", str(DATASET_INFO["classes"])),
        ("训练集路径", DATASET_INFO["train_path"]),
        ("验证集路径", DATASET_INFO["val_path"]),
        ("测试集路径", DATASET_INFO["test_path"]),
        ("授权协议", DATASET_INFO["license"]),
    ]
    add_formatted_table(doc, ["配置项", "值"], ds_rows, [4, 10.5])

    doc.add_page_break()

    # 第五章
    add_heading_styled("第五章  环境验证", 1)

    add_heading_styled("5.1 验证步骤", 2)
    add_body("完成全部安装后，按以下步骤逐项验证环境配置的正确性：")

    verify_steps = [
        ("1", "nvidia-smi", "确认GPU驱动正常，显示GPU型号与驱动版本"),
        ("2", "conda activate AIDetection", "确认虚拟环境可正常激活"),
        ("3", "python --version", "确认Python版本为3.7"),
        ("4", "python -c \"import torch; print(torch.cuda.is_available())\"", "确认PyTorch可检测到GPU"),
        ("5", "python -c \"from ultralytics import YOLO; print('OK')\"", "确认ultralytics框架安装成功"),
        ("6", "python -c \"import cv2; print(cv2.__version__)\"", "确认OpenCV安装成功"),
        ("7", "jupyter lab --version", "确认JupyterLab安装成功"),
    ]
    add_formatted_table(doc, ["序号", "验证命令", "预期结果"], verify_steps, [1.5, 7, 6])

    doc.add_paragraph("")

    add_heading_styled("5.2 验证结果", 2)
    add_body("所有验证步骤执行通过后，环境配置视为合格。如任一验证步骤失败，请参照对应章节重新执行安装操作。")

    # 第六章
    add_heading_styled("第六章  附录", 1)
    add_body("附录A：项目模型层目录结构")
    dir_lines = [
        "yolo_recognition_model/",
        "├── requirements.txt              # 模型层依赖清单",
        "├── environment.yml               # Conda环境配置文件",
        "├── dataset/",
        "│   └── data.yaml                 # 数据集配置",
        "├── training/",
        "│   └── configs/",
        "│       ├── baseline_config.yaml  # 基线训练配置",
        "│       ├── fast_iteration_config.yaml  # 快速训练配置",
        "│       ├── optimized_config.yaml       # 优化训练配置",
        "│       └── high_precision_config.yaml  # 高精度训练配置",
        "└── recognition-service/",
        "    ├── requirements.txt          # 服务层依赖清单",
        "    └── app/core/yolo_engine.py   # YOLO推理引擎",
    ]
    for line in dir_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # 保存
    filepath = f"{OUTPUT_DIR}\\环境安装报告_传统企业风格.docx"
    doc.save(filepath)
    print(f"[OK] 报告一已生成: {filepath}")


# ============================================================
# 报告二：现代简约风格
# ============================================================
def generate_report_2():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = Pt(18)

    ACCENT = RGBColor(0x00, 0x7A, 0xCC)
    DARK = RGBColor(0x2D, 0x2D, 0x2D)
    GRAY = RGBColor(0x75, 0x75, 0x75)
    ACCENT_HEX = "007ACC"

    def add_sec_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        # 添加色块标记
        run = p.add_run("■  ")
        run.font.color.rgb = ACCENT
        run.font.size = Pt(14)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # 分割线
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        run = p2.add_run("─" * 70)
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
        run.font.size = Pt(8)

    def add_sub_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = ACCENT
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_text(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.paragraph_format.space_after = Pt(4)

    def add_cmd(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"> {text}")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    def add_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_kv_table(data, col1="项目", col2="配置"):
        table = doc.add_table(rows=len(data) + 1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        # Header
        for i, h in enumerate([col1, col2]):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            set_cell_shading(cell, ACCENT_HEX)

        for r_idx, (k, v) in enumerate(data):
            cell_k = table.rows[r_idx + 1].cells[0]
            cell_v = table.rows[r_idx + 1].cells[1]
            cell_k.width = Cm(4)
            cell_v.width = Cm(11)

            cell_k.text = ""
            p = cell_k.paragraphs[0]
            run = p.add_run(k)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            set_cell_shading(cell_k, "EBF5FB")

            cell_v.text = ""
            p = cell_v.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        doc.add_paragraph("")
        return table

    # ---- 标题区 ----
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ENVIRONMENT SETUP REPORT")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("模型层环境安装报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = DARK
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("YOLO11 GPU Detection Environment  |  ")
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    run = p.add_run(REPORT_DATE_SHORT)
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(8)

    # ---- 1. 环境概览 ----
    add_sec_title("环境概览")
    add_kv_table([
        ("环境名称", "AIDetection"),
        ("Python版本", "3.7"),
        ("管理工具", "Conda（完整版，含env path）"),
        ("模型框架", "YOLO11（ultralytics >= 8.0.0）"),
        ("GPU加速", "CUDA >= 12.1 + cuDNN >= 8.9"),
        ("推荐IDE", "PyCharm Professional"),
        ("交互环境", "JupyterLab >= 4.0"),
    ])

    # ---- 2. 硬件与软件要求 ----
    add_sec_title("硬件与软件要求")

    add_sub_title("硬件要求")
    add_bullet("CPU：Intel Core i5+ / AMD Ryzen 5+")
    add_bullet("内存：16GB+（推荐32GB）")
    add_bullet("硬盘：SSD，可用空间 >= 50GB")
    add_bullet("GPU：NVIDIA，显存 >= 4GB（推荐 >= 8GB）")

    add_sub_title("软件要求")
    add_bullet("操作系统：Windows 10/11 64位 或 Ubuntu 20.04+")
    add_bullet("NVIDIA驱动 >= 525.60")
    add_bullet("CUDA Toolkit >= 12.1")
    add_bullet("cuDNN >= 8.9")

    # ---- 3. 安装流程 ----
    add_sec_title("安装流程")

    add_sub_title("Step 1  GPU驱动安装")
    add_text("从NVIDIA官网下载对应型号驱动，安装后执行验证：")
    add_cmd("nvidia-smi")
    add_text("确认输出中包含GPU型号、驱动版本及CUDA版本。")

    add_sub_title("Step 2  Conda安装")
    add_text("下载Anaconda3完整版安装包，安装时注意配置环境变量。验证：")
    add_cmd("conda --version")

    add_sub_title("Step 3  创建虚拟环境")
    add_cmd("conda create -n AIDetection python=3.7")
    add_cmd("conda activate AIDetection")
    add_text("环境路径：Anaconda安装目录\\envs\\AIDetection")

    add_sub_title("Step 4  安装依赖包")
    add_text("核心依赖（一行命令安装）：")
    add_cmd("pip install matplotlib numpy pandas Pillow PyYAML scikit-learn scipy torch torchaudio torchvision tqdm")
    add_text("")
    add_text("模型层扩展依赖：")
    add_cmd("pip install ultralytics>=8.0.0 opencv-python>=4.8.0 imagehash>=4.3.0 seaborn>=0.12.0 loguru>=0.7.0")

    # 依赖包明细表
    add_sub_title("依赖包明细")
    all_pkgs = CONDA_PACKAGES + EXTRA_MODEL_PACKAGES
    pkg_rows = [(name, ver, desc) for name, ver, desc in all_pkgs]
    table = doc.add_table(rows=len(pkg_rows) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(["包名", "版本", "用途"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, ACCENT_HEX)

    for r_idx, (name, ver, desc) in enumerate(pkg_rows):
        for c_idx, val in enumerate([name, ver, desc]):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    doc.add_paragraph("")

    add_sub_title("Step 5  开发工具配置")
    add_text("PyCharm（推荐）：File → Settings → Project → Python Interpreter → Add Conda Environment → 选择 AIDetection。")
    add_text("Jupyter：")
    add_cmd("pip install jupyterlab ipykernel")
    add_cmd("python -m ipykernel install --user --name AIDetection --display-name \"Python (AIDetection)\"")
    add_cmd("jupyter lab")

    # ---- 4. 模型层配置 ----
    add_sec_title("模型层配置")

    add_sub_title("YOLO11 模型规格")
    model_data = [
        ("yolo11n.pt", "Nano - 最快速度，适合快速验证"),
        ("yolo11s.pt", "Small - 精度与速度平衡，推荐默认使用"),
        ("yolo11m.pt", "Medium - 最高精度，适合离线推理"),
    ]
    add_kv_table(model_data, "模型", "规格 / 特点")

    add_sub_title("训练配置方案")
    table = doc.add_table(rows=len(YOLO_CONFIGS) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    headers = ["配置", "模型", "Epochs", "Batch", "ImgSz", "场景"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, ACCENT_HEX)

    for r_idx, cfg in enumerate(YOLO_CONFIGS):
        for c_idx, val in enumerate(cfg):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph("")

    add_sub_title("数据集信息")
    add_kv_table([
        ("来源", DATASET_INFO["source"]),
        ("类别数", str(DATASET_INFO["classes"])),
        ("训练集", DATASET_INFO["train_path"]),
        ("验证集", DATASET_INFO["val_path"]),
        ("测试集", DATASET_INFO["test_path"]),
        ("协议", DATASET_INFO["license"]),
    ])

    # ---- 5. 环境验证 ----
    add_sec_title("环境验证清单")

    verify_items = [
        ("nvidia-smi", "GPU驱动正常"),
        ("conda activate AIDetection", "环境激活成功"),
        ("python --version", "Python 3.7"),
        ("python -c \"import torch; print(torch.cuda.is_available())\"", "GPU可用 = True"),
        ("python -c \"from ultralytics import YOLO; print('OK')\"", "ultralytics安装成功"),
        ("python -c \"import cv2; print(cv2.__version__)\"", "OpenCV安装成功"),
        ("jupyter lab --version", "JupyterLab安装成功"),
    ]
    table = doc.add_table(rows=len(verify_items) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(["#", "验证命令", "预期结果"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, ACCENT_HEX)

    for r_idx, (cmd, expected) in enumerate(verify_items):
        cells = table.rows[r_idx + 1].cells
        cells[0].text = str(r_idx + 1)
        cells[1].text = cmd
        cells[2].text = expected
        for c in cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 保存
    filepath = f"{OUTPUT_DIR}\\环境安装报告_现代简约风格.docx"
    doc.save(filepath)
    print(f"[OK] 报告二已生成: {filepath}")


# ============================================================
# 报告三：技术文档风格
# ============================================================
def generate_report_3():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = Pt(16)

    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    BLUE = RGBColor(0x00, 0x52, 0x8A)
    CODE_BG = "F0F0F0"

    def add_doc_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_meta_line(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Consolas"

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = BLUE
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        # 下划线
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        run = p2.add_run("─" * 80)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(6)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"## {text}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"### {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_para(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p.paragraph_format.space_after = Pt(4)

    def add_code_block(lines):
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            # 背景色通过段落底纹
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
            p._element.get_or_add_pPr().append(shading)
        doc.add_paragraph("")

    def add_note(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"NOTE: {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x8B, 0x4C, 0x00)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.italic = True

    def add_spec_table(headers, rows):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Consolas"
            set_cell_shading(cell, "333333")

        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = "Consolas"
                if r_idx % 2 == 1:
                    set_cell_shading(cell, "F8F8F8")

        doc.add_paragraph("")
        return table

    # ---- 文档头部 ----
    add_doc_heading("模型层环境安装配置报告")
    add_meta_line(f"Document ID  : ENV-MODEL-2026-003")
    add_meta_line(f"Date         : {REPORT_DATE_SHORT}")
    add_meta_line(f"Status       : RELEASED")
    add_meta_line(f"Scope        : YOLO11 GPU Detection Environment")
    add_meta_line(f"Python       : 3.7 | Conda: AIDetection")

    p = doc.add_paragraph()
    run = p.add_run("=" * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(6)

    # ---- 1. 环境规格 ----
    add_h1("1. 环境规格")

    add_h2("1.1 基础环境")
    add_spec_table(["参数", "值"], [
        ("环境名称", "AIDetection"),
        ("Python", "3.7"),
        ("包管理器", "Conda (完整版, env path)"),
        ("模型框架", "YOLO11 (ultralytics >= 8.0.0)"),
        ("GPU加速", "CUDA >= 12.1 / cuDNN >= 8.9"),
        ("推荐IDE", "PyCharm Professional"),
        ("交互环境", "JupyterLab >= 4.0"),
    ])

    add_h2("1.2 硬件要求")
    add_spec_table(["组件", "最低配置", "推荐配置"], [
        ("CPU", "i5 / Ryzen 5", "i7 / Ryzen 7"),
        ("RAM", "16 GB", "32 GB"),
        ("Disk", "50 GB SSD", "100 GB NVMe SSD"),
        ("GPU", "4 GB VRAM", "8 GB+ VRAM"),
    ])

    # ---- 2. 安装步骤 ----
    add_h1("2. 安装步骤")

    add_h2("2.1 GPU驱动")
    add_code_block([
        "# 验证GPU驱动状态",
        "$ nvidia-smi",
        "",
        "# 预期输出包含:",
        "#   GPU型号 (如 NVIDIA GeForce RTX 3060)",
        "#   Driver Version >= 525.60",
        "#   CUDA Version >= 12.1",
    ])

    add_h2("2.2 Conda安装")
    add_para("下载Anaconda3完整版安装包，安装时配置环境变量。")
    add_code_block([
        "# 验证Conda安装",
        "$ conda --version",
        "$ conda info",
    ])

    add_h2("2.3 虚拟环境")
    add_code_block([
        "# 创建环境",
        "$ conda create -n AIDetection python=3.7",
        "",
        "# 激活环境",
        "$ conda activate AIDetection",
        "",
        "# 环境路径",
        "# <Anaconda_Install_Dir>/envs/AIDetection",
        "",
        "# 验证Python版本",
        "(AIDetection) $ python --version",
        "# Python 3.7.x",
    ])

    add_h2("2.4 依赖安装")
    add_h3("2.4.1 核心依赖")
    add_code_block([
        "$ pip install \\",
        "    matplotlib \\",
        "    numpy \\",
        "    pandas \\",
        "    Pillow \\",
        "    PyYAML \\",
        "    scikit-learn \\",
        "    scipy \\",
        "    torch \\",
        "    torchaudio \\",
        "    torchvision \\",
        "    tqdm",
    ])

    add_h3("2.4.2 模型层扩展依赖")
    add_code_block([
        "$ pip install \\",
        "    ultralytics>=8.0.0 \\",
        "    opencv-python>=4.8.0 \\",
        "    imagehash>=4.3.0 \\",
        "    seaborn>=0.12.0 \\",
        "    loguru>=0.7.0",
    ])

    add_h3("2.4.3 完整依赖清单")
    all_pkgs = CONDA_PACKAGES + EXTRA_MODEL_PACKAGES
    add_spec_table(
        ["Package", "Version", "Description"],
        [(n, v, d) for n, v, d in all_pkgs]
    )

    add_h2("2.5 开发工具")
    add_h3("PyCharm（推荐）")
    add_code_block([
        "# 配置路径",
        "File → Settings → Project → Python Interpreter",
        "  → Add Interpreter → Conda Environment",
        "  → Existing environment",
        "  → Interpreter: <Anaconda>/envs/AIDetection/python.exe",
    ])

    add_h3("Jupyter")
    add_code_block([
        "$ pip install jupyterlab>=4.0 ipykernel",
        "$ python -m ipykernel install --user \\",
        "    --name AIDetection \\",
        "    --display-name \"Python (AIDetection)\"",
        "$ jupyter lab",
    ])

    # ---- 3. 模型层配置 ----
    add_h1("3. 模型层配置")

    add_h2("3.1 YOLO11模型")
    add_para("本项目使用Ultralytics YOLO11架构，提供三种规格模型以适应不同场景：")
    add_spec_table(["Model", "Params", "Use Case"], [
        ("yolo11n.pt", "Nano (最小)", "快速验证、实时推理"),
        ("yolo11s.pt", "Small (平衡)", "默认训练配置（推荐）"),
        ("yolo11m.pt", "Medium (最大)", "高精度离线推理"),
    ])

    add_h2("3.2 训练配置")
    add_para("项目包含四套训练配置，通过YAML文件管理：")
    add_spec_table(
        ["Config", "Model", "Epochs", "Batch", "ImgSz", "Strategy"],
        [(c[0], c[1], c[2], c[3], c[4], c[5]) for c in YOLO_CONFIGS]
    )

    add_h3("关键训练参数")
    add_spec_table(["Parameter", "Value", "Note"], [
        ("optimizer", "AdamW", "自适应学习率优化器"),
        ("lr0", "0.001", "初始学习率"),
        ("lrf", "0.01", "终态学习率系数"),
        ("cos_lr", "True", "余弦退火调度"),
        ("amp", "True", "自动混合精度"),
        ("mosaic", "0.8~1.0", "Mosaic增强"),
        ("mixup", "0.0~0.2", "MixUp增强"),
        ("label_smoothing", "0.0~0.1", "标签平滑"),
        ("patience", "15~30", "早停耐心值"),
    ])

    add_h2("3.3 数据集")
    add_code_block([
        "# data.yaml",
        "path: .",
        f"train: {DATASET_INFO['train_path']}",
        f"val: {DATASET_INFO['val_path']}",
        f"test: {DATASET_INFO['test_path']}",
        f"nc: {DATASET_INFO['classes']}",
        f"source: {DATASET_INFO['source']}",
        f"license: {DATASET_INFO['license']}",
    ])

    # ---- 4. 验证 ----
    add_h1("4. 环境验证")
    add_code_block([
        "# 1. GPU驱动",
        "$ nvidia-smi",
        "",
        "# 2. 环境激活",
        "$ conda activate AIDetection",
        "",
        "# 3. Python版本",
        "$ python --version",
        "# => Python 3.7.x",
        "",
        "# 4. PyTorch GPU",
        "$ python -c \"import torch; print(torch.cuda.is_available())\"",
        "# => True",
        "",
        "# 5. ultralytics",
        "$ python -c \"from ultralytics import YOLO; print('OK')\"",
        "# => OK",
        "",
        "# 6. OpenCV",
        "$ python -c \"import cv2; print(cv2.__version__)\"",
        "",
        "# 7. Jupyter",
        "$ jupyter lab --version",
    ])

    add_para("全部验证通过后，环境配置视为合格。")

    # ---- 5. 项目结构 ----
    add_h1("5. 项目模型层目录结构")
    add_code_block([
        "yolo_recognition_model/",
        "├── requirements.txt                    # pip依赖清单",
        "├── environment.yml                     # Conda环境定义",
        "├── dataset/",
        "│   └── data.yaml                       # 数据集配置",
        "├── training/",
        "│   ├── configs/",
        "│   │   ├── baseline_config.yaml        # 基线配置",
        "│   │   ├── fast_iteration_config.yaml  # 快速迭代",
        "│   │   ├── optimized_config.yaml       # 优化配置 ★",
        "│   │   └── high_precision_config.yaml  # 高精度配置",
        "│   └── scripts/",
        "│       └── train_yolov11_optimized.py  # 训练入口脚本",
        "├── recognition-service/",
        "│   ├── requirements.txt                # 服务层依赖",
        "│   └── app/core/yolo_engine.py         # YOLO推理引擎",
        "└── yolo_training_outputs/",
        "    └── runs/train/args.yaml            # 训练运行参数记录",
    ])

    # 保存
    filepath = f"{OUTPUT_DIR}\\环境安装报告_技术文档风格.docx"
    doc.save(filepath)
    print(f"[OK] 报告三已生成: {filepath}")


# ============================================================
# 主入口
# ============================================================
if __name__ == "__main__":
    import os
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_report_1()
    generate_report_2()
    generate_report_3()
    print("\n全部报告生成完毕。")
